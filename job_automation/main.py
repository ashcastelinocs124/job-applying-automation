import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import agentops
from agents import Agents
from tasks import Tasks
from crewai import Crew, Agent, Task
from sample import (
    sample_resume,
    sample_job_posting,
    sample_company_culture,
)
from utils import parse_markdown_to_pdf
from docx import Document


tracer = agentops.start_trace(
    trace_name="CrewAI Job Posting",
    tags=["crew-job-posting-example", "agentops-example"],
)
tasks = Tasks()
agents = Agents()
company_description = (
    "We are a software company that builds AI-powered tools for businesses."
)
company_domain = "https://www.agentops.ai"
hiring_needs = "We are looking for a software engineer with 3 years of experience in Python and Django."
specific_benefits = "We offer a competitive salary, health insurance, and a 401k plan."
company_name = "agentops"


researcher_agent = agents.research_agent()
writer_agent = agents.writer_agent()
review_agent = agents.review_agent()
cover_letter_agent = agents.cover_letter_agent()
resume_agent = agents.resume_agent()

research_company_culture_task = tasks.research_company_culture_task(
    researcher_agent, company_description, company_domain
)
research_role_requirements_task = tasks.research_role_requirements_task(
    researcher_agent, hiring_needs
)
review_and_edit_job_posting_task = tasks.research_company_background(
    researcher_agent, company_name
)
generate_cover_letter = tasks.generate_cover_letter_task(
    cover_letter_agent, sample_resume, sample_job_posting, sample_company_culture
)
generate_resume = tasks.generate_resume(
    resume_agent, sample_resume, sample_job_posting, sample_company_culture
)


# Instantiate the crew with a sequential process
crew = Crew(
    agents=[
        researcher_agent,
        writer_agent,
        review_agent,
        cover_letter_agent,
        resume_agent,
    ],
    tasks=[
        research_company_culture_task,
        research_role_requirements_task,
        review_and_edit_job_posting_task,
        generate_cover_letter,
        generate_resume,
    ],
)

result = crew.kickoff()

print("=== JOB APPLICATION PACKAGE ===")
print(result)

# Convert result to string for parsing
full_output = str(result)

# Parse the output to extract cover letter and resume
cover_letter_content = ""
resume_content = ""

# Simple parsing logic - look for task outputs
lines = full_output.split("\n")
current_section = None

for line in lines:
    line_lower = line.lower()

    # Detect section headers
    if any(keyword in line_lower for keyword in ["cover letter", "cover letter:"]):
        current_section = "cover_letter"
        continue
    elif any(keyword in line_lower for keyword in ["resume", "resume:"]):
        current_section = "resume"
        continue
    elif any(
        keyword in line_lower
        for keyword in ["research", "company culture", "role requirements"]
    ):
        current_section = "research"
        continue

    # Collect content based on current section
    if current_section == "cover_letter" and line.strip():
        cover_letter_content += line + "\n"
    elif current_section == "resume" and line.strip():
        resume_content += line + "\n"

# If parsing failed, use the full output for both
if not cover_letter_content.strip() and not resume_content.strip():
    # Fallback: split the output roughly in half
    lines = full_output.split("\n")
    mid_point = len(lines) // 2
    cover_letter_content = "\n".join(lines[:mid_point])
    resume_content = "\n".join(lines[mid_point:])

# Generate PDF for resume
if resume_content.strip():
    pdf_filename = "tailored_resume.pdf"
    try:
        parse_markdown_to_pdf(resume_content, pdf_filename)
        print(f"\n✅ Resume saved as: {pdf_filename}")
    except Exception as e:
        print(f"\n⚠️  Error generating PDF: {e}")
        # Save as text file as fallback
        with open("tailored_resume.txt", "w") as f:
            f.write(resume_content)
        print("📄 Resume saved as: tailored_resume.txt (fallback)")
else:
    print("\n⚠️  No resume content found to save")

# Generate DOC for cover letter
if cover_letter_content.strip():
    doc_filename = "cover_letter.docx"
    try:
        doc = Document()
        doc.add_heading("Cover Letter", 0)
        doc.add_paragraph(cover_letter_content)
        doc.save(doc_filename)
        print(f"✅ Cover letter saved as: {doc_filename}")
    except Exception as e:
        print(f"⚠️  Error generating DOC: {e}")
        # Save as text file as fallback
        with open("cover_letter.txt", "w") as f:
            f.write(cover_letter_content)
        print("📄 Cover letter saved as: cover_letter.txt (fallback)")
else:
    print("⚠️  No cover letter content found to save")

print("\n=== Files Generated ===")
print("📄 tailored_resume.pdf (or .txt fallback)")
print("📄 cover_letter.docx (or .txt fallback)")
